from pathlib import Path
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


TITLE_COLOR = RGBColor(31, 78, 121)
SUBTITLE_COLOR = RGBColor(47, 117, 181)
HEADING_COLOR = RGBColor(31, 78, 121)


def add_centered_run(paragraph, text, size_pt, bold=False, color=None, italic=False):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color


def add_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = HEADING_COLOR


def add_divider(doc):
    p = doc.add_paragraph("—" * 72)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def parse_sections(text):
    sections = []
    section_pattern = re.compile(
        r"\{\s*section:\s*\"([^\"]+)\"\s*,\s*qa:\s*\[(.*?)\]\s*\}",
        re.DOTALL,
    )
    qa_pattern = re.compile(
        r"q:\s*\"([^\"]+)\"\s*,\s*a:\s*`([^`]+)`",
        re.DOTALL,
    )

    for section_match in section_pattern.finditer(text):
        section_title = section_match.group(1).strip()
        qa_block = section_match.group(2)
        qa_list = []
        for qa_match in qa_pattern.finditer(qa_block):
            question = qa_match.group(1).strip()
            answer = qa_match.group(2).strip()
            qa_list.append((question, answer))
        sections.append((section_title, qa_list))

    return sections


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(10.5)


def add_answer_paragraphs(doc, answer_text):
    lines = answer_text.splitlines()
    for line in lines:
        if not line.strip():
            doc.add_paragraph("")
            continue
        doc.add_paragraph(line)


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    source = root / "Resume_questions.md"
    target = root / "Resume_questions_styled.docx"

    text = source.read_text(encoding="utf-8")
    sections = parse_sections(text)

    doc = Document()

    # Header
    add_centered_run(doc.add_paragraph(), "JAVA FULL STACK DEVELOPER", 24, bold=True, color=TITLE_COLOR)
    add_centered_run(doc.add_paragraph(), "Interview Preparation Guide", 14, bold=True, color=SUBTITLE_COLOR)
    add_centered_run(doc.add_paragraph(), "Prepared for: Nadipi Srujan Reddy", 11)
    add_centered_run(doc.add_paragraph(), "B.Tech CSE | CGPA: 8.5 | Hyderabad, India", 10.5)
    add_centered_run(
        doc.add_paragraph(),
        "100 Comprehensive Interview Questions | From a Senior Java Full Stack Developer (10+ Years)",
        9.5,
        italic=True,
    )

    add_divider(doc)

    # How to use section
    add_heading(doc, "How to Use This Guide")
    add_bullets(
        doc,
        [
            "Answer each question out loud as if in a real interview — time yourself (aim for 2-4 mins per question).",
            "Use the answer area below each question to jot down key points, code snippets, or diagrams.",
            "For coding/design questions, draw architecture diagrams on paper while explaining.",
            "Relate every answer back to your actual projects — interviewers value real-world experience over textbook answers.",
            "Sections are ordered from fundamentals to advanced — complete them in order for best results.",
        ],
    )

    doc.add_paragraph("")

    # Questions and answers
    question_number = 1
    for section_title, qa_list in sections:
        add_heading(doc, section_title)
        doc.add_paragraph("")
        for question, answer in qa_list:
            q_text = f"Q{question_number}. {question}"
            q_paragraph = doc.add_paragraph()
            q_run = q_paragraph.add_run(q_text)
            q_run.bold = True
            q_run.font.size = Pt(11)

            add_answer_paragraphs(doc, answer)
            doc.add_paragraph("")
            question_number += 1

    doc.save(target)
    print(f"Wrote: {target}")


if __name__ == "__main__":
    main()
